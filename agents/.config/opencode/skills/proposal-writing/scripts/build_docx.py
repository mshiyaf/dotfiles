#!/usr/bin/env python3
"""Build a proposal DOCX from Markdown using the proposal Word template."""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS_DIR = ROOT / "assets"
TEMPLATE_PATH = ROOT / "templates" / "proposal-template.docx"


@dataclass(frozen=True)
class Company:
    legal_name: str
    address: str
    email: str
    phone: str
    website: str
    logo: str
    footer_badge: str | None = None


COMPANIES = {
    "sprdh": Company(
        legal_name="SPRDH Solutions Private Limited",
        address="First Floor, Kerala Startup Mission, Kerala Technology Innovation Zone, Kalamassery, Kochi, Kerala 683503, India",
        email="hello@sprdh.com",
        phone="+91 799 428 66 60",
        website="www.sprdh.com",
        logo="sprdh-logo.jpeg",
        footer_badge="sprdh-anniversary-badge.png",
    ),
    "texol": Company(
        legal_name="Texol - iCohort Business Ventures Pvt. Ltd",
        address="Pantheerankavu, Kozhikode, Kerala, India, 673019",
        email="support@texolworld.com",
        phone="+91 90379 81865",
        website="texolworld.com",
        logo="texol-logo.png",
    ),
}


FRONTMATTER_RE = re.compile(r"\A---\s*\n(.*?)\n---\s*\n?", re.DOTALL)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a styled proposal DOCX from Markdown.")
    parser.add_argument("draft", type=Path, help="Markdown proposal draft")
    parser.add_argument("-o", "--output", type=Path, required=True, help="Output .docx path")
    parser.add_argument("--company", help="Sender identity: sprdh or texol")
    parser.add_argument("--title", help="Proposal title for the cover page")
    parser.add_argument("--client", help="Client name for the cover page")
    parser.add_argument("--date", dest="proposal_date", help="Proposal date")
    parser.add_argument("--document-id", help="Document ID shown in the header and cover details")
    parser.add_argument("--version", default="1.0", help="Proposal version")
    parser.add_argument("--author", default="Shiyaf C", help="Author shown on the cover page")
    parser.add_argument("--approver", default="Saalim Bin Ali", help="Approver shown on the cover page")
    parser.add_argument("--client-logo", type=Path, help="Optional client logo for the Submitted To block")
    return parser.parse_args()


def clean_markdown_text(text: str) -> str:
    text = re.sub(r"<([^>]+)>", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.strip()


def infer_metadata(markdown: str) -> dict[str, str]:
    metadata: dict[str, str] = {}
    for raw in markdown.splitlines():
        line = raw.strip().rstrip(" ")
        if not line:
            continue
        title = re.match(r"^#\s+(.+)$", line)
        if title and "title" not in metadata:
            metadata["title"] = clean_markdown_text(title.group(1))
            continue
        prepared_for = re.match(r"^\*\*Prepared for:\*\*\s*(.+?)(?:\s{2,})?$", line, re.IGNORECASE)
        if prepared_for and "client" not in metadata:
            metadata["client"] = clean_markdown_text(prepared_for.group(1))
            continue
        proposal_date = re.match(r"^\*\*Date:\*\*\s*(.+?)(?:\s{2,})?$", line, re.IGNORECASE)
        if proposal_date and "date" not in metadata:
            metadata["date"] = clean_markdown_text(proposal_date.group(1))
            continue
    return metadata


def read_draft(path: Path) -> tuple[dict[str, str], str]:
    text = path.read_text(encoding="utf-8")
    match = FRONTMATTER_RE.match(text)
    if not match:
        return infer_metadata(text), text

    metadata: dict[str, str] = {}
    for line in match.group(1).splitlines():
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        metadata[key.strip().lower()] = value.strip().strip('"\'')
    markdown = text[match.end() :]
    inferred = infer_metadata(markdown)
    inferred.update(metadata)
    return inferred, markdown


def first_present(*values: str | None, default: str = "") -> str:
    for value in values:
        if value:
            return value
    return default


def acronym(value: str) -> str:
    uppercase = re.findall(r"\b[A-Z0-9]{3,}\b", value)
    if uppercase:
        return uppercase[0][:8]
    words = re.findall(r"[A-Za-z0-9]+", value)
    if not words:
        return "CLIENT"
    if len(words) == 1:
        return words[0][:8].upper()
    return "".join(word[0] for word in words[:6]).upper()


def compact_date(value: str) -> str:
    for fmt in ("%d %B %Y", "%d %b %Y", "%Y-%m-%d", "%d/%m/%Y"):
        try:
            return datetime.strptime(value, fmt).strftime("%d%m%y")
        except ValueError:
            pass
    return date.today().strftime("%d%m%y")


def default_document_id(company_key: str, client: str, proposal_date: str, version: str) -> str:
    prefix = "SP" if company_key == "sprdh" else "TX"
    return f"{prefix}/{acronym(client)}/{compact_date(proposal_date)}/TCP/{version}"


def clear_block(block) -> None:
    element = block._element
    for child in list(element):
        element.remove(child)


def clear_document_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def normalize_template_fonts(document: Document) -> None:
    for style in document.styles:
        element = style._element
        for fonts in element.findall(".//w:rFonts", namespaces=element.nsmap):
            for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
                key = qn(f"w:{attr}")
                value = fonts.get(key)
                if value and value.startswith("Montserrat"):
                    fonts.set(key, "Montserrat")


def set_cell_text(cell, text: str, bold: bool = False, font_size_pt: float | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.bold = bold
    if font_size_pt:
        run.font.size = Pt(font_size_pt)


def style_table(table) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, left: int = 90, right: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        margin = tc_mar.find(qn(f"w:{edge}"))
        if margin is None:
            margin = OxmlElement(f"w:{edge}")
            tc_mar.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def set_cell_nowrap(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = OxmlElement("w:noWrap")
    tc_pr.append(no_wrap)


def set_cell_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_style = OxmlElement("w:rStyle")
    run_style.set(qn("w:val"), "Hyperlink")
    run_properties.append(run_style)
    run.append(run_properties)

    run_text = OxmlElement("w:t")
    run_text.text = text
    run.append(run_text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_runs_from_inline(paragraph, text: str) -> None:
    token_re = re.compile(r"(\[[^\]]+\]\((?:https?://|mailto:)[^)]+\)|<(?:(?:https?://|mailto:)[^>\s]+)>|\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        markdown_link = re.fullmatch(r"\[([^\]]+)\]\(((?:https?://|mailto:)[^)]+)\)", token)
        if markdown_link:
            add_hyperlink(paragraph, markdown_link.group(1), markdown_link.group(2))
        elif token.startswith("<"):
            url = token[1:-1]
            add_hyperlink(paragraph, url, url)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def set_body_paragraph_spacing(paragraph, before: float = 0, after: float = 8, line: float | None = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    if line is not None:
        paragraph.paragraph_format.line_spacing = line


def set_hanging_list_indent(paragraph, level: int = 0) -> None:
    left = 0.35 + (level * 0.25)
    paragraph.paragraph_format.left_indent = Inches(left)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1


def set_paragraph_outline_level(paragraph, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    outline_level = p_pr.first_child_found_in("w:outlineLvl")
    if outline_level is None:
        outline_level = OxmlElement("w:outlineLvl")
        p_pr.append(outline_level)
    outline_level.set(qn("w:val"), str(level - 1))


def add_styled_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    set_body_paragraph_spacing(
        paragraph,
        before=12 if level == 1 else 8,
        after=6 if level == 1 else 4,
        line=1.0,
    )
    set_paragraph_outline_level(paragraph, level)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Montserrat"
    run.font.color.rgb = RGBColor(55, 96, 146)
    run.font.size = Pt(15 if level == 1 else 13)


def add_page_field(paragraph) -> None:
    # python-docx has no page-number API, so the Word PAGE field is raw OOXML.
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    fld.append(run)
    paragraph._p.append(fld)


def mark_fields_for_update(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_field(paragraph, instruction: str, placeholder: str) -> None:
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run.append(begin)
    paragraph._p.append(begin_run)

    instruction_run = OxmlElement("w:r")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction_text.text = instruction
    instruction_run.append(instruction_text)
    paragraph._p.append(instruction_run)

    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    paragraph._p.append(separate_run)

    paragraph.add_run(placeholder)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph._p.append(end_run)


def add_table_of_contents(document: Document) -> None:
    heading = document.add_paragraph()
    set_body_paragraph_spacing(heading, before=12, after=12, line=1.0)
    run = heading.add_run("Table of Contents")
    run.bold = True
    run.font.name = "Montserrat"
    run.font.color.rgb = RGBColor(55, 96, 146)
    run.font.size = Pt(15)

    toc = document.add_paragraph()
    add_field(toc, 'TOC \\o "1-2" \\h \\z \\u', "Update table of contents in Word or LibreOffice.")
    document.add_page_break()


def configure_header_footer(document: Document, company_key: str, company: Company, document_id: str) -> None:
    section = document.sections[0]
    section.different_first_page_header_footer = False

    header = section.header
    clear_block(header)
    paragraph = header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_path = ASSETS_DIR / company.logo
    if not logo_path.exists():
        raise FileNotFoundError(f"Missing logo asset for {company_key}: {logo_path}")
    paragraph.add_run().add_picture(str(logo_path), width=Inches(1.85))
    paragraph.add_run("\t" + document_id)

    footer = section.footer
    clear_block(footer)
    page_paragraph = footer.add_paragraph()
    page_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if company.footer_badge:
        badge_path = ASSETS_DIR / company.footer_badge
        if badge_path.exists():
            page_paragraph.add_run().add_picture(str(badge_path), width=Inches(0.45))
            page_paragraph.add_run("  ")
    page_paragraph.add_run("Page | ")
    add_page_field(page_paragraph)

    contact = footer.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(f"{company.website} | {company.email}")


def add_cover(document: Document, metadata: dict[str, str], company_key: str, company: Company, args: argparse.Namespace) -> None:
    title = first_present(args.title, metadata.get("title"), default="Techno-Commercial Proposal")
    client = first_present(args.client, metadata.get("client"), default="Client")
    proposal_date = first_present(args.proposal_date, metadata.get("date"), default=date.today().strftime("%d %B %Y"))
    version = first_present(args.version, metadata.get("version"), default="1.0")
    document_id = first_present(
        args.document_id,
        metadata.get("document_id"),
        metadata.get("document-id"),
        default=default_document_id(company_key, title, proposal_date, version),
    )

    for _ in range(3):
        document.add_paragraph("", style="NS Title")

    title_paragraph = document.add_paragraph(style="NS Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True

    document.add_paragraph("")
    submitted = document.add_paragraph("Submitted To")
    submitted.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if args.client_logo:
        if not args.client_logo.exists():
            raise FileNotFoundError(f"Client logo not found: {args.client_logo}")
        logo_paragraph = document.add_paragraph(style="NS Title")
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_paragraph.add_run().add_picture(str(args.client_logo), width=Inches(1.5))

    client_paragraph = document.add_paragraph(client)
    client_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    client_paragraph.runs[0].bold = True

    document.add_paragraph("")
    details = document.add_table(rows=1, cols=4)
    style_table(details)
    for cell, width in zip(details.rows[0].cells, [1.2, 3.1, 1.0, 0.6]):
        set_cell_width(cell, width)
    set_cell_text(details.cell(0, 0), "Document ID:", bold=True)
    set_cell_text(details.cell(0, 1), document_id, font_size_pt=7)
    set_cell_nowrap(details.cell(0, 1))
    set_cell_text(details.cell(0, 2), "Version:", bold=True)
    set_cell_text(details.cell(0, 3), version)

    document.add_paragraph("")
    approvals = document.add_table(rows=3, cols=4)
    style_table(approvals)
    for idx, value in enumerate(["", "Name", "Signature", "Date"]):
        set_cell_text(approvals.cell(0, idx), value, bold=True)
    set_cell_text(approvals.cell(1, 0), "Author:", bold=True)
    set_cell_text(approvals.cell(1, 1), args.author)
    set_cell_text(approvals.cell(2, 0), "Approver", bold=True)
    set_cell_text(approvals.cell(2, 1), args.approver)
    set_cell_text(approvals.cell(2, 3), proposal_date)

    document.add_paragraph("")
    document.add_paragraph("")
    document.add_paragraph("Amendment Record", style="NS Section Heading 1")
    amendments = document.add_table(rows=2, cols=5)
    style_table(amendments)
    for idx, value in enumerate(["Version", "Date", "Description", "Amended By", "Approved By"]):
        set_cell_text(amendments.cell(0, idx), value, bold=True)
    for idx, value in enumerate([version, proposal_date, "Initial proposal", args.author, args.approver]):
        set_cell_text(amendments.cell(1, idx), value)

    document.add_page_break()

    for _ in range(4):
        document.add_paragraph("")
    document.add_paragraph(
        f"Copyright (c) {date.today().year} by {company.legal_name}. All rights reserved. "
        "This document is confidential and may be used only to evaluate this proposal. "
        "It may not be reproduced, published, or shared with any third party without written permission."
    )
    document.add_paragraph("Proposal Expiration Date: 1 Month from date of submission (30 Calendar Days)", style="Style Bold Line spacing:  1.5 lines")
    document.add_paragraph("")
    document.add_paragraph("For any questions, clarifications, or information, please reach out to:")
    document.add_paragraph(company.legal_name)
    document.add_paragraph(f"Email: {company.email}")
    document.add_paragraph(f"Phone: {company.phone}")
    document.add_page_break()


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and "|" in lines[index]:
        if not is_table_separator(lines[index]):
            rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
        index += 1
    return rows, index


def markdown_table_column_widths(column_count: int) -> list[float]:
    if column_count == 2:
        return [2.1, 4.4]
    if column_count == 3:
        return [0.9, 4.15, 1.45]
    if column_count == 4:
        return [1.1, 2.6, 1.4, 1.4]
    usable_width = 6.5
    return [usable_width / column_count] * column_count


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    style_table(table)
    column_widths = markdown_table_column_widths(width)
    repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(rows):
        table_row = table.rows[row_idx]
        table_row.height = Pt(22 if row_idx == 0 else 20)
        table_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for col_idx in range(width):
            text = row[col_idx] if col_idx < len(row) else ""
            cell = table.cell(row_idx, col_idx)
            set_cell_width(cell, column_widths[col_idx])
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            add_runs_from_inline(paragraph, text)
            if row_idx == 0:
                set_cell_shading(cell, "376092")
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(9)
            else:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            set_body_paragraph_spacing(paragraph, after=0, line=1.0)
    spacer = document.add_paragraph()
    set_body_paragraph_spacing(spacer, before=0, after=4, line=0.4)


def normalize_heading_text(text: str) -> str:
    return re.sub(r"^\d+(?:\.\d+)*[.)]?\s+", "", text).strip()


def body_lines_without_cover_preamble(markdown: str) -> list[str]:
    lines = markdown.splitlines()
    for index, line in enumerate(lines):
        if re.match(r"^##\s+", line.strip()):
            return lines[index:]
    return lines


def add_markdown_body(document: Document, markdown: str, cover_title: str) -> None:
    lines = body_lines_without_cover_preamble(markdown)
    index = 0
    current_paragraph = None
    heading_counter = 0
    subheading_counter = 0

    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped:
            current_paragraph = None
            index += 1
            continue

        if "|" in stripped and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            rows, index = parse_table(lines, index)
            add_markdown_table(document, rows)
            current_paragraph = None
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = normalize_heading_text(heading.group(2).strip())
            if not (level == 1 and text == cover_title):
                if level <= 2:
                    heading_counter += 1
                    subheading_counter = 0
                    add_styled_heading(document, f"{heading_counter}.  {text}", 1)
                else:
                    subheading_counter += 1
                    add_styled_heading(document, f"{heading_counter}.{subheading_counter}.  {text}", 2)
            current_paragraph = None
            index += 1
            continue

        bullet = re.match(r"^(\s*)[-*+]\s+(.*)$", raw)
        if bullet:
            paragraph = document.add_paragraph()
            level = 1 if len(bullet.group(1)) >= 2 else 0
            set_hanging_list_indent(paragraph, level)
            paragraph.add_run("•\t")
            add_runs_from_inline(paragraph, bullet.group(2).strip())
            current_paragraph = None
            index += 1
            continue

        numbered = re.match(r"^(\s*)\d+[.)]\s+(.*)$", raw)
        if numbered:
            paragraph = document.add_paragraph()
            level = 1 if len(numbered.group(1)) >= 2 else 0
            set_hanging_list_indent(paragraph, level)
            paragraph.add_run(numbered.group(0).strip().split(maxsplit=1)[0] + "\t")
            add_runs_from_inline(paragraph, numbered.group(2).strip())
            current_paragraph = None
            index += 1
            continue

        if current_paragraph is None:
            current_paragraph = document.add_paragraph()
            set_body_paragraph_spacing(current_paragraph)
            add_runs_from_inline(current_paragraph, stripped)
        else:
            current_paragraph.add_run(" ")
            add_runs_from_inline(current_paragraph, stripped)
        index += 1


def main() -> int:
    args = parse_args()
    metadata, markdown = read_draft(args.draft)
    company_key = first_present(args.company, metadata.get("company")).lower()
    if company_key not in COMPANIES:
        valid = ", ".join(sorted(COMPANIES))
        print(f"error: company must resolve to one of: {valid}. Use --company or frontmatter company.", file=sys.stderr)
        return 2
    if not TEMPLATE_PATH.exists():
        print(f"error: missing template: {TEMPLATE_PATH}", file=sys.stderr)
        return 2

    company = COMPANIES[company_key]
    title = first_present(args.title, metadata.get("title"), default="Techno-Commercial Proposal")
    client = first_present(args.client, metadata.get("client"), default="Client")
    proposal_date = first_present(args.proposal_date, metadata.get("date"), default=date.today().strftime("%d %B %Y"))
    version = first_present(args.version, metadata.get("version"), default="1.0")
    document_id = first_present(
        args.document_id,
        metadata.get("document_id"),
        metadata.get("document-id"),
        default=default_document_id(company_key, title, proposal_date, version),
    )

    document = Document(TEMPLATE_PATH)
    normalize_template_fonts(document)
    clear_document_body(document)
    try:
        configure_header_footer(document, company_key, company, document_id)
    except FileNotFoundError as exc:
        print(f"error: {exc}", file=sys.stderr)
        return 2
    add_cover(document, metadata, company_key, company, args)
    add_table_of_contents(document)
    add_markdown_body(document, markdown, title)
    mark_fields_for_update(document)
    document.core_properties.title = title
    document.core_properties.author = company.legal_name

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
